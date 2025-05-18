import os
import argparse
import logging
from docx import Document
from extractor import GeminiRecipeExtractor
from sql_writer import generate_sql_inserts
from config import DEFAULT_OUTPUT_DIR
from logging_config import setup_logging

def extract_tables_with_headers(docx_path):
    doc = Document(docx_path)
    tables = []
    paragraphs = list(doc.paragraphs)
    para_idx = 0
    for table in doc.tables:
        header = None
        for i in range(para_idx-1, -1, -1):
            text = paragraphs[i].text.strip()
            if text:
                header = text
                para_idx = i + 1
                break
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        tables.append({'header': header, 'rows': rows})
        para_idx += 1
    return tables

def table_to_markdown(table):
    rows = table['rows']
    if not rows:
        return ''
    header_row = '| ' + ' | '.join(rows[0]) + ' |'
    sep_row = '| ' + ' | '.join(['---'] * len(rows[0])) + ' |'
    data_rows = ['| ' + ' | '.join(r) + ' |' for r in rows[1:]]
    return '\n'.join([header_row, sep_row] + data_rows)

def main():
    parser = argparse.ArgumentParser(description='Enterprise Recipe Uploader (Pure Gemini Table Mode)')
    parser.add_argument('--input-dir', required=True, help='Directory containing recipe .docx files')
    parser.add_argument('--output-dir', default=DEFAULT_OUTPUT_DIR, help='Directory to write SQL and logs')
    args = parser.parse_args()

    setup_logging(args.output_dir)
    extractor = GeminiRecipeExtractor()
    recipe_files = [f for f in os.listdir(args.input_dir) if f.endswith('.docx')]
    all_sql_statements = []
    for recipe_file in recipe_files:
        logging.info(f"Processing {recipe_file}...")
        docx_path = os.path.join(args.input_dir, recipe_file)
        tables = extract_tables_with_headers(docx_path)
        # Build a single markdown-rich prompt
        markdown_sections = []
        component_counter = 1
        for table in tables[:-1]:  # Disregard the last table (portion options)
            component_name = table['header'] or f"{{Recipe Name}} - Component {component_counter}"
            component_counter += 1
            markdown_table = table_to_markdown(table)
            markdown_sections.append(f"Component: {component_name}\n{markdown_table}")
        markdown_prompt = '\n\n'.join(markdown_sections)
        prompt = f"""This recipe contains the following tables:\n\n{markdown_prompt}\n\nInstructions:\n- For each component, extract all ingredients with name, quantity (numeric), unit, and calories.\n- If a quantity cell contains both raw and cooked weights (e.g., '250g after cooking 640g'), extract both as 'raw_quantity' and 'cooked_quantity'.\n- If a table has a 'Total' row, use its quantity as the component's base_quantity. If both before and after cooking are present, extract both.\n- Always extract numeric values for each ingredient's quantity and unit.\n- If the component name is hard to identify, use '{{Recipe Name}} - Component {component_counter}' where {component_counter} is a number.\n- Return a JSON object with:\n  - name: Recipe name\n  - components: list of components, each with:\n    - name: component name\n    - base_quantity: (from the Total row, if present)\n    - ingredients: list of ingredients (name, quantity, unit, calories, raw_quantity, cooked_quantity as available)\n"""
        recipe_struct = extractor.extract_recipe(prompt)
        # Ensure recipe_struct has a 'name' field
        if recipe_struct is not None and 'name' not in recipe_struct:
            recipe_struct['name'] = os.path.splitext(recipe_file)[0]
        if recipe_struct and 'components' in recipe_struct:
            sql_statements = generate_sql_inserts(recipe_struct)
            all_sql_statements.extend(sql_statements)
        else:
            logging.error(f"Failed to extract structured recipe for {recipe_file}")
    os.makedirs(args.output_dir, exist_ok=True)
    output_file = os.path.join(args.output_dir, 'recipe_inserts.sql')
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write("-- Generated SQL statements for recipe import\n")
        f.write("BEGIN;\n\n")
        for statement in all_sql_statements:
            f.write(statement + "\n")
        f.write("\nCOMMIT;\n")
    logging.info(f"SQL statements written to {output_file}")

if __name__ == "__main__":
    main() 