#!/usr/bin/env python3
import sys
import json
import os
import re
from pathlib import Path
from docx import Document
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def extract_tables_with_headers_and_positions(docx_path):
    """Extract tables and their associated headers with position information from a Word document."""
    try:
        doc = Document(docx_path)
        tables = []
        
        # Get all tables first
        for table_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            
            tables.append({
                'rows': rows,
                'table_index': table_idx
            })
        
        return tables
    except Exception as e:
        logger.error(f"Error reading document: {e}")
        return []

def extract_document_content(docx_path):
    """Extract all text content from the Word document."""
    try:
        doc = Document(docx_path)
        content = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text)
        
        # Extract table content
        for table in doc.tables:
            content.append("")  # Add spacing
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    content.append(row_text)
            content.append("")  # Add spacing
        
        return "\n".join(content)
    except Exception as e:
        logger.error(f"Error extracting document content: {e}")
        return "Error extracting document content"

def detect_recipe_boundaries(docx_path):
    """Detect recipe boundaries in the document by looking for recipe titles."""
    try:
        doc = Document(docx_path)
        recipe_starts = []
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            # Look for recipe title patterns
            if (text and 
                (text.isupper() or 
                 any(keyword in text.lower() for keyword in ['recipe', 'meal', 'dish', 'preparation']) or
                 re.match(r'^[A-Z][A-Z\s]+$', text)) and
                len(text) > 3 and len(text) < 100):
                recipe_starts.append(i)
        
        return recipe_starts
    except Exception as e:
        logger.error(f"Error detecting recipe boundaries: {e}")
        return []

def extract_recipe_sections(docx_path):
    """Extract different recipe sections from the document with enhanced boundary detection."""
    try:
        doc = Document(docx_path)
        recipe_starts = []
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            # Enhanced recipe title detection
            if (text and 
                (text.isupper() or 
                 any(keyword in text.lower() for keyword in ['recipe', 'meal', 'dish', 'preparation', 'soup', 'salad', 'main', 'appetizer', 'dessert']) or
                 re.match(r'^[A-Z][A-Z\s]+$', text) or
                 re.match(r'^[A-Z][a-zA-Z\s]+Recipe', text) or
                 re.match(r'^[A-Z][a-zA-Z\s]+Soup', text) or
                 re.match(r'^[A-Z][a-zA-Z\s]+Salad', text)) and
                len(text) > 3 and len(text) < 100 and
                not any(keyword in text.lower() for keyword in ['ingredient', 'component', 'total', 'portion', 'serving'])):
                recipe_starts.append(i)
        
        # If no clear boundaries found, try alternative detection
        if not recipe_starts:
            # Look for patterns like "Recipe 1:", "Recipe 2:", etc.
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if re.match(r'^Recipe\s+\d+', text, re.IGNORECASE):
                    recipe_starts.append(i)
        
        sections = []
        for i, start_idx in enumerate(recipe_starts):
            end_idx = recipe_starts[i + 1] if i + 1 < len(recipe_starts) else len(doc.paragraphs)
            recipe_name = doc.paragraphs[start_idx].text.strip()
            sections.append({
                'start': start_idx,
                'end': end_idx,
                'name': recipe_name,
                'index': i
            })
        
        return sections
    except Exception as e:
        logger.error(f"Error extracting recipe sections: {e}")
        return []

def assign_tables_to_recipes(tables, recipe_sections):
    """Assign tables to recipes based on their content and position."""
    # This is a simplified approach - in practice you might need more sophisticated logic
    # For now, we'll distribute tables evenly among recipes
    
    if not recipe_sections:
        return {0: tables}  # All tables go to first recipe
    
    tables_per_recipe = len(tables) // len(recipe_sections)
    remainder = len(tables) % len(recipe_sections)
    
    recipe_tables = {}
    table_idx = 0
    
    for i, section in enumerate(recipe_sections):
        # Calculate how many tables this recipe should get
        num_tables = tables_per_recipe
        if i < remainder:
            num_tables += 1
        
        recipe_tables[i] = tables[table_idx:table_idx + num_tables]
        table_idx += num_tables
    
    return recipe_tables

def table_to_markdown(table):
    """Convert a table to markdown format."""
    if not table['rows']:
        return ""
    
    markdown = "| " + " | ".join(table['rows'][0]) + " |\n"
    markdown += "| " + " | ".join(["---"] * len(table['rows'][0])) + " |\n"
    
    for row in table['rows'][1:]:
        markdown += "| " + " | ".join(row) + " |\n"
    
    return markdown

def parse_quantity(quantity_str):
    """Parse quantity string to numeric value."""
    if not quantity_str:
        return 0
    
    try:
        # Remove common units and extract number
        import re
        quantity_str = str(quantity_str).strip()
        # Extract first number found
        match = re.search(r'(\d+(?:\.\d+)?)', quantity_str)
        if match:
            return float(match.group(1))
        return 0
    except:
        return 0

def parse_nutrition_value(nutrition_str):
    """Parse nutrition value from string."""
    if not nutrition_str:
        return 0
    
    try:
        import re
        nutrition_str = str(nutrition_str).strip()
        # Extract first number found
        match = re.search(r'(\d+(?:\.\d+)?)', nutrition_str)
        if match:
            return float(match.group(1))
        return 0
    except:
        return 0

def is_summary_row(ingredient):
    """Check if ingredient row is a summary/total row."""
    name = ingredient.get('name', '').lower()
    return any(keyword in name for keyword in ['total', 'sum', 'subtotal'])

def get_ingredient_name(ingredient):
    """Extract ingredient name from ingredient data."""
    return ingredient.get('name', '').strip()

def extract_recipe_from_section(doc, section, tables_for_recipe):
    """Extract recipe data from a specific section of the document."""
    try:
        recipe = {
            "name": section['name'],
            "description": f"Extracted from {section['name']}",
            "packaging": "",
            "objective": "",
            "itemCode": "",
            "components": []
        }
        
        # First pass: Look for general information table (usually the first table)
        if tables_for_recipe and len(tables_for_recipe) > 0:
            first_table = tables_for_recipe[0]
            if first_table['rows']:
                # Check if this looks like a general info table (2 columns, key-value pairs)
                if len(first_table['rows']) > 0 and len(first_table['rows'][0]) == 2:
                    for row in first_table['rows']:
                        if len(row) == 2:
                            key = row[0].strip().lower()
                            value = row[1].strip()
                            
                            if 'package' in key:
                                recipe["packaging"] = value
                            elif 'objective' in key:
                                recipe["objective"] = value
                            elif 'item code' in key or 'itemcode' in key:
                                recipe["itemCode"] = value
                            elif 'item name' in key or 'itemname' in key:
                                # Update recipe name if it's more specific
                                if value and value != recipe["name"]:
                                    recipe["name"] = value
        
        # Second pass: Process ingredient tables (skip the first table if it was general info)
        ingredient_tables = tables_for_recipe[1:] if len(tables_for_recipe) > 1 and recipe["packaging"] else tables_for_recipe
        
        # Get component names from document paragraphs that appear before each table
        component_names = extract_component_names_from_document(doc, len(ingredient_tables))
        
        for i, table in enumerate(ingredient_tables):
            # Use extracted component name or fallback to generic name
            component_name = component_names[i] if i < len(component_names) and component_names[i] else f"Component {i+1}"
            
            component = {
                "name": component_name,
                "ingredients": []
            }
            
            if table['rows']:
                # Assume first row is header
                headers = table['rows'][0]
                for row in table['rows'][1:]:
                    if len(row) >= 2:  # At least name and quantity
                        ingredient_name = row[0].strip()
                        if ingredient_name and not is_summary_row({"name": ingredient_name}):
                            quantity = parse_quantity(row[1] if len(row) > 1 else "0")
                            unit = row[2] if len(row) > 2 else "g"
                            
                            # Try to extract nutrition values from various columns
                            calories = 0
                            fat = 0
                            protein = 0
                            carbohydrates = 0
                            
                            # Look for nutrition data in remaining columns
                            for col_idx in range(3, min(len(row), 7)):  # Check columns 3-6
                                cell_value = row[col_idx] if col_idx < len(row) else ""
                                cell_lower = cell_value.lower()
                                
                                # Try to identify nutrition type and extract value
                                if any(keyword in cell_lower for keyword in ['cal', 'kcal', 'calories']):
                                    calories = parse_nutrition_value(cell_value)
                                elif any(keyword in cell_lower for keyword in ['fat', 'lipids']):
                                    fat = parse_nutrition_value(cell_value)
                                elif any(keyword in cell_lower for keyword in ['protein', 'prot']):
                                    protein = parse_nutrition_value(cell_value)
                                elif any(keyword in cell_lower for keyword in ['carb', 'carbs', 'carbohydrates']):
                                    carbohydrates = parse_nutrition_value(cell_value)
                                elif col_idx == 3:  # Default to calories if it's the 4th column
                                    calories = parse_nutrition_value(cell_value)
                            
                            ingredient = {
                                "name": ingredient_name,
                                "quantity": quantity,
                                "unit": unit,
                                "calories": calories,
                                "fat": fat,
                                "protein": protein,
                                "carbohydrates": carbohydrates
                            }
                            component["ingredients"].append(ingredient)
            
            recipe["components"].append(component)
        
        return recipe
        
    except Exception as e:
        logger.error(f"Error extracting recipe from section: {e}")
        return None

def extract_component_names_from_document(doc, num_tables):
    """Extract component names from document paragraphs that appear before tables."""
    try:
        component_names = []
        table_count = 0
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            
            # Skip empty paragraphs
            if not text:
                continue
            
            # Skip metadata text (contains numbers and units)
            if re.search(r'\d+\s*(g|kcal|cal|gram|pound|oz|ml|l)', text.lower()):
                continue
            
            # Skip summary/total text
            if any(keyword in text.lower() for keyword in ['total', 'summary', 'yield', 'calories per gram', 'cooked yield']):
                continue
            
            # Pattern 1: Numbered sections like "3. Mushroom Sautéed with Gochujang"
            numbered_match = re.match(r'^\d+\.\s*(.+)', text)
            if numbered_match:
                component_name = numbered_match.group(1).strip()
                if component_name and len(component_name) > 2:
                    # Clean up the component name
                    name = component_name.strip()
                    # Remove common suffixes
                    name = re.sub(r'\s*ingredients?\s*$', '', name, flags=re.IGNORECASE)
                    name = re.sub(r'\s*composition\s*$', '', name, flags=re.IGNORECASE)
                    name = re.sub(r'\s*breakdown\s*$', '', name, flags=re.IGNORECASE)
                    
                    if name and len(name) > 2:
                        component_names.append(name)
                        table_count += 1
                        continue
            
            # Pattern 2: Bold section titles like "Beef Bulgogi Ingredients"
            # Check if this looks like a component title (contains keywords)
            if any(keyword in text.lower() for keyword in [
                'ingredients', 'composition', 'sautéed', 'dressing', 'bulgogi', 
                'mushroom', 'sauce', 'marinade', 'breakdown'
            ]):
                # Clean up the component name
                name = text.strip()
                # Remove common suffixes
                name = re.sub(r'\s*ingredients?\s*$', '', name, flags=re.IGNORECASE)
                name = re.sub(r'\s*composition\s*$', '', name, flags=re.IGNORECASE)
                name = re.sub(r'\s*breakdown\s*$', '', name, flags=re.IGNORECASE)
                
                if name and len(name) > 2:
                    component_names.append(name)
                    table_count += 1
                    continue
            
            # Pattern 3: Check if this paragraph is followed by a table
            # This is a fallback for when we can't identify the pattern clearly
            if table_count < num_tables and len(text) > 3 and len(text) < 100:
                # Check if this looks like a component title (not too long, not too short)
                if not any(keyword in text.lower() for keyword in ['total', 'summary', 'note', 'instruction', 'yield', 'calories', 'cooked']):
                    # Additional check: should not contain measurement patterns
                    if not re.search(r'\d+\s*(g|kcal|cal|gram|pound|oz|ml|l)', text):
                        component_names.append(text)
                        table_count += 1
        
        # If we didn't find enough component names, fill with generic names
        while len(component_names) < num_tables:
            component_names.append(f"Component {len(component_names) + 1}")
        
        # Return only the number of names we need
        return component_names[:num_tables]
        
    except Exception as e:
        logger.error(f"Error extracting component names from document: {e}")
        # Return generic names as fallback
        return [f"Component {i+1}" for i in range(num_tables)]

def extract_multiple_recipes_from_docx(docx_path):
    """Extract multiple recipes from a Word document with proper table association."""
    try:
        doc = Document(docx_path)
        tables = extract_tables_with_headers_and_positions(docx_path)
        document_content = extract_document_content(docx_path)
        
        logger.info(f"Found {len(tables)} tables in document")
        
        # Try to detect recipe sections
        recipe_sections = extract_recipe_sections(docx_path)
        logger.info(f"Detected {len(recipe_sections)} recipe sections")
        for i, section in enumerate(recipe_sections):
            logger.info(f"Section {i}: '{section['name']}' (para {section['start']}-{section['end']})")
        
        recipes = []
        
        if len(recipe_sections) > 1:
            # Multiple recipes detected - assign tables to recipes
            recipe_tables = assign_tables_to_recipes(tables, recipe_sections)
            
            for i, section in enumerate(recipe_sections):
                logger.info(f"Processing section {i}: {section['name']}")
                tables_for_recipe = recipe_tables.get(i, [])
                logger.info(f"  Assigned {len(tables_for_recipe)} tables to recipe {i}")
                
                recipe = extract_recipe_from_section(doc, section, tables_for_recipe)
                if recipe and recipe['components']:
                    logger.info(f"Extracted recipe '{recipe['name']}' with {len(recipe['components'])} components")
                    recipes.append(recipe)
                else:
                    logger.warning(f"No valid recipe extracted from section '{section['name']}'")
        else:
            # Single recipe or no clear boundaries - use original logic
            logger.info("Using single recipe extraction logic")
            recipe = extract_recipe_from_docx_single(docx_path)
            if recipe:
                recipes.append(recipe)
        
        # If no recipes found, create a default one
        if not recipes:
            logger.info("No recipes found, creating default recipe")
            recipe = {
                "name": Path(docx_path).stem,
                "description": f"Extracted from {Path(docx_path).name}",
                "packaging": "",
                "objective": "",
                "itemCode": "",
                "components": []
            }
            
            # First pass: Check if first table contains general information
            if tables and len(tables) > 0:
                first_table = tables[0]
                if first_table['rows']:
                    # Check if this looks like a general info table (2 columns, key-value pairs)
                    if len(first_table['rows']) > 0 and len(first_table['rows'][0]) == 2:
                        for row in first_table['rows']:
                            if len(row) == 2:
                                key = row[0].strip().lower()
                                value = row[1].strip()
                                
                                if 'package' in key:
                                    recipe["packaging"] = value
                                elif 'objective' in key:
                                    recipe["objective"] = value
                                elif 'item code' in key or 'itemcode' in key:
                                    recipe["itemCode"] = value
                                elif 'item name' in key or 'itemname' in key:
                                    # Update recipe name if it's more specific
                                    if value and value != recipe["name"]:
                                        recipe["name"] = value
            
            # Second pass: Process ingredient tables (skip the first table if it was general info)
            ingredient_tables = tables[1:] if len(tables) > 1 and recipe["packaging"] else tables[:-1]  # Skip the last table (usually portion options)
            
            # Get component names from document paragraphs that appear before each table
            component_names = extract_component_names_from_document(doc, len(ingredient_tables))
            
            for i, table in enumerate(ingredient_tables):
                # Use extracted component name or fallback to generic name
                component_name = component_names[i] if i < len(component_names) and component_names[i] else f"Component {i+1}"
                
                component = {
                    "name": component_name,
                    "ingredients": []
                }
                
                if table['rows']:
                    # Assume first row is header
                    headers = table['rows'][0]
                    for row in table['rows'][1:]:
                        if len(row) >= 2:  # At least name and quantity
                            ingredient_name = row[0].strip()
                            if ingredient_name and not is_summary_row({"name": ingredient_name}):
                                quantity = parse_quantity(row[1] if len(row) > 1 else "0")
                                unit = row[2] if len(row) > 2 else "g"
                                
                                # Try to extract nutrition values from various columns
                                calories = 0
                                fat = 0
                                protein = 0
                                carbohydrates = 0
                                
                                # Look for nutrition data in remaining columns
                                for col_idx in range(3, min(len(row), 7)):  # Check columns 3-6
                                    cell_value = row[col_idx] if col_idx < len(row) else ""
                                    cell_lower = cell_value.lower()
                                    
                                    # Try to identify nutrition type and extract value
                                    if any(keyword in cell_lower for keyword in ['cal', 'kcal', 'calories']):
                                        calories = parse_nutrition_value(cell_value)
                                    elif any(keyword in cell_lower for keyword in ['fat', 'lipids']):
                                        fat = parse_nutrition_value(cell_value)
                                    elif any(keyword in cell_lower for keyword in ['protein', 'prot']):
                                        protein = parse_nutrition_value(cell_value)
                                    elif any(keyword in cell_lower for keyword in ['carb', 'carbs', 'carbohydrates']):
                                        carbohydrates = parse_nutrition_value(cell_value)
                                    elif col_idx == 3:  # Default to calories if it's the 4th column
                                        calories = parse_nutrition_value(cell_value)
                                
                                ingredient = {
                                    "name": ingredient_name,
                                    "quantity": quantity,
                                    "unit": unit,
                                    "calories": calories,
                                    "fat": fat,
                                    "protein": protein,
                                    "carbohydrates": carbohydrates
                                }
                                component["ingredients"].append(ingredient)
                
                recipe["components"].append(component)
            
            recipes.append(recipe)
        
        logger.info(f"Final result: {len(recipes)} recipes extracted")
        return recipes, document_content
        
    except Exception as e:
        logger.error(f"Error extracting multiple recipes: {e}")
        return [], "Error extracting document content"

def extract_recipe_from_docx_single(docx_path):
    """Extract single recipe data from a Word document (original function)."""
    try:
        tables = extract_tables_with_headers_and_positions(docx_path)
        
        if not tables:
            return None
        
        recipe = {
            "name": Path(docx_path).stem,
            "description": f"Extracted from {Path(docx_path).name}",
            "packaging": "",
            "objective": "",
            "itemCode": "",
            "components": []
        }
        
        # First pass: Check if first table contains general information
        if tables and len(tables) > 0:
            first_table = tables[0]
            if first_table['rows']:
                # Check if this looks like a general info table (2 columns, key-value pairs)
                if len(first_table['rows']) > 0 and len(first_table['rows'][0]) == 2:
                    for row in first_table['rows']:
                        if len(row) == 2:
                            key = row[0].strip().lower()
                            value = row[1].strip()
                            
                            if 'package' in key:
                                recipe["packaging"] = value
                            elif 'objective' in key:
                                recipe["objective"] = value
                            elif 'item code' in key or 'itemcode' in key:
                                recipe["itemCode"] = value
                            elif 'item name' in key or 'itemname' in key:
                                # Update recipe name if it's more specific
                                if value and value != recipe["name"]:
                                    recipe["name"] = value
        
        # Second pass: Process ingredient tables (skip the first table if it was general info)
        ingredient_tables = tables[1:] if len(tables) > 1 and recipe["packaging"] else tables
        
        # Get component names from document paragraphs that appear before each table
        doc = Document(docx_path)
        component_names = extract_component_names_from_document(doc, len(ingredient_tables))
        
        for i, table in enumerate(ingredient_tables):
            # Use extracted component name or fallback to generic name
            component_name = component_names[i] if i < len(component_names) and component_names[i] else f"Component {i+1}"
            
            component = {
                "name": component_name,
                "ingredients": []
            }
            
            if table['rows']:
                # Assume first row is header
                headers = table['rows'][0]
                for row in table['rows'][1:]:
                    if len(row) >= 2:  # At least name and quantity
                        ingredient_name = row[0].strip()
                        if ingredient_name and not is_summary_row({"name": ingredient_name}):
                            quantity = parse_quantity(row[1] if len(row) > 1 else "0")
                            unit = row[2] if len(row) > 2 else "g"
                            
                            # Try to extract nutrition values from various columns
                            calories = 0
                            fat = 0
                            protein = 0
                            carbohydrates = 0
                            
                            # Look for nutrition data in remaining columns
                            for col_idx in range(3, min(len(row), 7)):  # Check columns 3-6
                                cell_value = row[col_idx] if col_idx < len(row) else ""
                                cell_lower = cell_value.lower()
                                
                                # Try to identify nutrition type and extract value
                                if any(keyword in cell_lower for keyword in ['cal', 'kcal', 'calories']):
                                    calories = parse_nutrition_value(cell_value)
                                elif any(keyword in cell_lower for keyword in ['fat', 'lipids']):
                                    fat = parse_nutrition_value(cell_value)
                                elif any(keyword in cell_lower for keyword in ['protein', 'prot']):
                                    protein = parse_nutrition_value(cell_value)
                                elif any(keyword in cell_lower for keyword in ['carb', 'carbs', 'carbohydrates']):
                                    carbohydrates = parse_nutrition_value(cell_value)
                                elif col_idx == 3:  # Default to calories if it's the 4th column
                                    calories = parse_nutrition_value(cell_value)
                            
                            ingredient = {
                                "name": ingredient_name,
                                "quantity": quantity,
                                "unit": unit,
                                "calories": calories,
                                "fat": fat,
                                "protein": protein,
                                "carbohydrates": carbohydrates
                            }
                            component["ingredients"].append(ingredient)
            
            recipe["components"].append(component)
        
        return recipe
        
    except Exception as e:
        logger.error(f"Error extracting recipe: {e}")
        return None

def main():
    """Main function to be called from Node.js."""
    if len(sys.argv) != 2:
        print(json.dumps({"error": "Usage: python extract_recipe.py <docx_path>"}))
        sys.exit(1)
    
    docx_path = sys.argv[1]
    
    if not os.path.exists(docx_path):
        print(json.dumps({"error": f"File not found: {docx_path}"}))
        sys.exit(1)
    
    try:
        recipes, document_content = extract_multiple_recipes_from_docx(docx_path)
        if recipes:
            print(json.dumps({
                "success": True, 
                "recipes": recipes,
                "documentContent": document_content
            }))
        else:
            print(json.dumps({
                "error": "Failed to extract recipe data",
                "documentContent": document_content
            }))
    except Exception as e:
        print(json.dumps({"error": str(e)}))

if __name__ == "__main__":
    main() 